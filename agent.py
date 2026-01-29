import json
import random
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
import matplotlib.pyplot as plt
import os

# ---------------------------
# 1️ LOAD JSON INPUT
# ---------------------------
json_file = "infopercept_analytics.json"

if not os.path.exists(json_file):
    raise FileNotFoundError(f"{json_file} not found in current directory!")

with open(json_file, "r") as f:
    analytics_data = json.load(f)

# ---------------------------
# 2️ GENERATE FAKE GRAPHS
# ---------------------------
months = list(range(1, 13))
incidents = [random.randint(900, 1800) for _ in months]
latencies = [random.randint(120, 220) for _ in months]
threat_score = [random.uniform(0.5, 0.95) for _ in months]

# Line chart: Incidents
plt.figure(figsize=(6,4))
plt.plot(months, incidents, marker='o', color='blue', linewidth=2)
plt.title("Monthly Security Incidents Detected")
plt.xlabel("Month")
plt.ylabel("Incidents")
plt.grid(True)
incidents_chart = "incidents.png"
plt.savefig(incidents_chart)
plt.close()

# Line chart: Latency
plt.figure(figsize=(6,4))
plt.plot(months, latencies, marker='s', color='red', linewidth=2)
plt.title("Average Model Latency (ms)")
plt.xlabel("Month")
plt.ylabel("Latency (ms)")
plt.grid(True)
latency_chart = "latency.png"
plt.savefig(latency_chart)
plt.close()

# Line chart: Threat Score
plt.figure(figsize=(6,4))
plt.plot(months, threat_score, marker='^', color='green', linewidth=2)
plt.title("Threat Detection Confidence Score")
plt.xlabel("Month")
plt.ylabel("Score")
plt.grid(True)
threat_chart = "threat_score.png"
plt.savefig(threat_chart)
plt.close()

# ---------------------------
# 3️  CREATE DOCX REPORT
# ---------------------------
doc = Document()

# ---- Title Page ----
doc.add_heading(analytics_data["company"], 0)
doc.add_heading(analytics_data["project_name"], level=1)
doc.add_paragraph(
    "This document presents a comprehensive AI-driven cybersecurity analytics report "
    "for INFOPERCEPT. The report includes system overview, analytics insights, "
    "visualizations, risk assessment, and recommendations."
).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

doc.add_page_break()

# ---- Add Sections ----
sections = [
    "Executive Summary",
    "System Architecture",
    "Data Collection & Ingestion",
    "Feature Engineering",
    "Model Training & Validation",
    "Threat Detection Results",
    "Behavioral Analytics Insights",
    "Regional Performance Analysis",
    "Latency & Scalability Analysis",
    "False Positive Reduction Strategy",
    "Operational Impact",
    "Risk Assessment & Mitigation",
    "Future Roadmap",
    "Conclusion"
]

for section in sections:
    doc.add_heading(section, level=2)
    for _ in range(4):
        para = doc.add_paragraph(
            "INFOPERCEPT leverages cutting-edge AI and ML technologies to process "
            "large-scale telemetry data and provide actionable security insights. "
            "This section elaborates on design choices, analytics, and outcomes."
        )
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_page_break()

# ---- Add Charts ----
doc.add_heading("Analytics Visualizations", level=2)
doc.add_paragraph("Key metrics and trends from the project are visualized below:")

doc.add_paragraph("Monthly Security Incidents:")
doc.add_picture(incidents_chart, width=Inches(5))
doc.add_paragraph("Average Model Latency (ms):")
doc.add_picture(latency_chart, width=Inches(5))
doc.add_paragraph("Threat Detection Confidence Score:")
doc.add_picture(threat_chart, width=Inches(5))

doc.add_page_break()

# ---- Add Fake Images (Dashboard placeholders) ----
# You can replace these with real screenshots
for i in range(1, 4):
    doc.add_heading(f"Dashboard Visualization {i}", level=2)
    doc.add_paragraph("This dashboard shows real-time threat intelligence metrics and alerts.")
    # Use the same chart as placeholder
    doc.add_picture(incidents_chart, width=Inches(5))
    doc.add_page_break()

# ---- Add Tables (Key Metrics) ----
doc.add_heading("Key Metrics Table", level=2)
table = doc.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Service'
hdr_cells[1].text = 'Requests/Day'
hdr_cells[2].text = 'Avg Latency(ms)'
hdr_cells[3].text = 'p99 Latency(ms)'
hdr_cells[4].text = 'False Positive Rate'

for svc in analytics_data["services"]:
    row_cells = table.add_row().cells
    row_cells[0].text = svc["name"]
    row_cells[1].text = str(svc["requests_per_day"])
    row_cells[2].text = str(svc["avg_latency_ms"])
    row_cells[3].text = str(svc["p99_latency_ms"])
    row_cells[4].text = str(svc["false_positive_rate"])

doc.add_page_break()

# ---- Appendix JSON ----
doc.add_heading("Appendix: Source Analytics JSON", level=2)
doc.add_paragraph(json.dumps(analytics_data, indent=2))

# ---- Save DOCX ----
doc_path = "INFOPERCEPT_AI_Styled_Report.docx"
doc.save(doc_path)

print(f"Report generated successfully: {doc_path}")
